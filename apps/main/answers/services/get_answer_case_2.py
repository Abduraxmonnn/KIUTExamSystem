# Python
from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
import psycopg2
import os
import json
import datetime

# Django
from django.db import connection as default_db

# Project
from apps.main.answers.models import Answer

base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..', '..'))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")

query = """SELECT students_student.student_id, students_student.full_name, student_groups_studentgroup.code, directions_direction.code,
            directions_direction.name, subjects_subject.full_name, subjects_subject.code, teachers_teacher.full_name, score,
            teachers_subjects_teachersubject.language, answer_json,exam_schedule_examschedule.exam_date, exam_schedule_examschedule.room from answers_answer
            join students_student on students_student.id = answers_answer.student_id 
            join student_groups_studentgroup on students_student.group_id = student_groups_studentgroup.id
            join teachers_subjects_teachersubject on teachers_subjects_teachersubject.group_id = student_groups_studentgroup.id and 
            answers_answer.subject_id = teachers_subjects_teachersubject.subject_id 
            join teachers_teacher on teachers_subjects_teachersubject.teacher_id = teachers_teacher.id
            join subjects_subject on answers_answer.subject_id = subjects_subject.id
            join directions_direction on subjects_subject.direction_id = directions_direction.id
            join exam_schedule_examschedule on subjects_subject.id=exam_schedule_examschedule.subject_id and 
            students_student.id = exam_schedule_examschedule.student_id
            where answers_answer.stage=2 and students_student.student_id = %s and subjects_subject.code = %s"""


def generate_and_save_docx_case_2(student_id: str = None, subject_code: str = None):
    host = default_db.settings_dict['HOST']
    user = default_db.settings_dict['USER']
    password = default_db.settings_dict['PASSWORD']
    database = default_db.settings_dict['NAME']
    connection = psycopg2.connect(host=host, database=database, user=user, password=password)
    connection.autocommit = True
    cursor = connection.cursor()

    cursor.execute(query, (student_id, subject_code))
    data = cursor.fetchone()

    if not data:
        return {
            'status': 'error',
            'message': 'connection to db is failed'
        }

    # [0] = student_id
    # [1] = student_name
    # [2] = group_code
    # [3] = direction_code
    # [4] = direction_name
    # [5] = subject_name
    # [6] = subject_code
    # [7] = teacher_name
    # [8] = score
    # [9] = language
    # [10] = json
    # [11] = exam_date_time
    # [12] = exam_room
    exam_date_time = data[11] - datetime.timedelta(minutes=120)
    text_info = f"""Student ID:  {data[0]}
    Student's Fullname:  {data[1]}
    Group:  {data[2]}
    Direction:  {data[3]}-{data[4]}
    Subject:  {data[6]}-{data[5]}
    Language:  {data[9]}
    Teacher's Fullname:  {data[7]}
    Exam date:  {exam_date_time.strftime("%d.%m.%Y")}
    Exam time:  {exam_date_time.strftime("%H:%M")}
    Exam room:  {data[12]}
    Correct answers:  {data[8]}
    Wrong answers:  {20 - data[8]}
    Created time of the document: {datetime.date.today().strftime("%B %d, %Y")}  //  {datetime.datetime.now().strftime("%H:%M:%S")}
    """

    gray_icon = os.path.join(base_dir, 'static/gray.jpg')
    green_icon = os.path.join(base_dir, 'static/green.jpg')
    red_icon = os.path.join(base_dir, 'static/red.jpg')
    rate_icon = os.path.join(base_dir, 'static/rate.png')

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    heading = doc.add_heading(text_info, level=0)
    infocursor = heading.add_run()
    infocursor.add_break()
    infocursor.add_picture(gray_icon)
    infocursor.add_text(" - student's answer")
    infocursor.add_break()
    infocursor.add_picture(green_icon)
    infocursor.add_text(" - correct answer")
    infocursor.add_break()
    infocursor.add_picture(red_icon)
    infocursor.add_text(" - wrong answer")
    infocursor.add_break()
    infocursor.add_picture(rate_icon)
    infocursor.add_text(" - questions' rate")
    infocursor.add_break()
    for run in heading.runs:
        run.font.size = Pt(18)
    doc.add_paragraph('TEST').alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph1 = doc.add_paragraph()
    cursor = paragraph1.add_run()
    cursor.add_break(WD_BREAK.LINE)

    json_string = json.dumps(data[10], indent=4, ensure_ascii=False)
    json_data = json.loads(json_string)
    list_options = ['option1', 'option2', 'option3', 'option4']
    count = 1
    for question_id, question_data in json_data.items():
        for question in question_data['question']:
            for i in range(question['rate']):
                cursor.add_picture(rate_icon, Inches(0.2))
            variant = 0
            cursor.bold = True
            cursor.add_text(f"{count}.")
            for value in question["question"]:
                if 'type' in value:
                    if value['type'] == 'text':
                        cursor.add_text(value['content'])
                    elif value['type'] == 'image':
                        pics_path = os.path.join(base_dir, str(value['content'][1:]))
                        cursor.add_picture(pics_path)
            cursor.add_break()
            cursor.bold = False
            for option_key in list_options:
                option = question[option_key]
                if variant == 0:
                    char = 'A)'
                elif variant == 1:
                    char = 'B)'
                elif variant == 2:
                    char = 'C)'
                else:
                    char = 'D)'
                cursor.add_text(char)
                is_correct = option[-1]['correct']
                if question_data['picked'] == option_key:
                    cursor.add_picture(gray_icon, Inches(0.2))
                if is_correct:
                    cursor.add_picture(green_icon, Inches(0.2))
                else:
                    cursor.add_picture(red_icon, Inches(0.2))
                for value in option:
                    if 'type' in value:
                        if value['type'] == 'text':
                            cursor.add_text(value['content'])
                        elif value['type'] == 'image':
                            pics_path = os.path.join(base_dir, str(value['content'][1:]))
                            cursor.add_picture(pics_path)
                variant += 1
                cursor.add_break()
        count += 1
        cursor.add_break()

    answer_obj = Answer.objects.filter(subject__code=subject_code, student__student_id=student_id,
                                       stage=2).last()

    output_dir = (os.path.join(base_dir, 'media', 'case_2', subject_code))
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir,
                               f"{answer_obj.question.language}_{answer_obj.subject.full_name}_{student_id}.docx")
    doc.save(str(output_path))

    relative_output_path = os.path.relpath(output_path, base_dir)
    response = {
        "student_id": student_id,
        "subject": answer_obj.subject.full_name,
        "score": 2,
        "answer": relative_output_path
    }
    return {
        'status': 'successfully',
        'message': response
    }
